"""Test case loader and DOCX report generator for compliance mapping.

This module provides:
1. Loading test cases from XLSX files
2. Matching discovered vulnerabilities to test cases
3. Generating DOCX format compliance reports with the specified template
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from openpyxl import Workbook, load_workbook
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


logger = logging.getLogger(__name__)


@dataclass
class TestCase:
    """Represents a test case from the XLSX file."""
    id: str
    name: str
    category: str = ""
    severity: str = ""
    description: str = ""
    remediation: str = ""
    status: str = "NOT_TESTED"  # NOT_TESTED, PASS, FAIL, PARTIAL
    matched_vulnerability_id: str | None = None
    evidence: str = ""
    url_endpoint: str = ""


@dataclass
class VulnerabilityMatch:
    """Represents a matched vulnerability to a test case."""
    test_case: TestCase
    vulnerability: dict[str, Any]
    confidence: float  # 0.0 to 1.0
    match_reason: str


def load_test_cases_from_xlsx(xlsx_path: str | Path) -> list[TestCase]:
    """Load test cases from an XLSX file.
    
    Supports two formats:
    
    **Simple Format **(6 columns)
    - ID: Test case identifier
    - Name: Test case name (e.g., "SQL 注入漏洞检查")
    - Category: Vulnerability category
    - Severity: Expected severity
    - Description: Test case description
    - Remediation: Remediation guidance
    
    **Extended Format **(12 columns - your format)
    - 序号：Serial number
    - 用例名称：Test case name (KEY COLUMN)
    - 类型：Type
    - 前提条件：Prerequisites
    - 测试步骤：Test steps
    - 预期结果：Expected results
    - 基础分类：Basic category
    - 威胁类型：Threat type
    - 备注：Notes
    - 版本：Version
    - 是否实现自动化：Automation status
    - 级别：Level (severity)
    
    The loader automatically detects the format and extracts the test case name
    as the primary matching key.
    """
    xlsx_path = Path(xlsx_path)
    if not xlsx_path.exists():
        raise FileNotFoundError(f"Test case file not found: {xlsx_path}")
    
    test_cases: list[TestCase] = []
    wb = load_workbook(xlsx_path, read_only=True, data_only=True)
    ws = wb.active
    
    if ws is None:
        raise ValueError("No active worksheet found in XLSX file")
    
    # Assume first row is header
    headers = [str(cell.value) if cell.value else "" for cell in next(ws.iter_rows(min_row=1, max_row=1))]
    
    # Find column indices - support both formats
    col_map = {name.lower().strip(): idx for idx, name in enumerate(headers)}
    
    # Detect format by checking for key columns
    has_name = "name" in col_map or "用例名称" in col_map
    has_severity = "severity" in col_map or "级别" in col_map
    has_category = "category" in col_map or "基础分类" in col_map or "类型" in col_map
    
    if not has_name:
        raise ValueError("XLSX must contain either 'Name' or '用例名称' column")
    
    for row_idx, row in enumerate(ws.iter_rows(min_row=2), start=2):
        values = [str(cell.value) if cell.value else "" for cell in row]
        
        def get_val(col_names: list[str], default: str = "") -> str:
            for col_name in col_names:
                idx = col_map.get(col_name.lower())
                if idx is not None and idx < len(values):
                    val = values[idx].strip()
                    if val:
                        return val
            return default
        
        # Get test case name (required)
        tc_name = get_val(["name", "用例名称"], "Unnamed Test Case")
        
        # Get other fields with fallbacks for different formats
        tc_id = get_val(["id", "序号"], f"TC-{row_idx}")
        tc_category = get_val(["category", "基础分类", "类型"], "")
        tc_severity = get_val(["severity", "级别"], "")
        tc_description = get_val(["description", "测试步骤", "预期结果"], "")
        tc_remediation = get_val(["remediation", "备注"], "")
        
        test_cases.append(TestCase(
            id=tc_id,
            name=tc_name,
            category=tc_category,
            severity=tc_severity,
            description=tc_description,
            remediation=tc_remediation,
        ))
    
    logger.info(f"Loaded {len(test_cases)} test cases from {xlsx_path}")
    return test_cases


def match_vulnerabilities_to_test_cases(
    test_cases: list[TestCase],
    vulnerabilities: list[dict[str, Any]]
) -> list[VulnerabilityMatch]:
    """Match discovered vulnerabilities to test cases using LLM-friendly heuristics.
    
    Matching strategy:
    1. Keyword matching on test case name vs vulnerability title/description
    2. CWE/CVE matching if available
    3. Endpoint/URL pattern matching
    """
    matches: list[VulnerabilityMatch] = []
    
    # Build keyword index for test cases
    tc_keywords: dict[str, list[str]] = {}
    for tc in test_cases:
        keywords = []
        # Extract keywords from name (e.g., "SQL 注入漏洞检查" -> ["SQL", "注入", "SQL 注入"])
        name_lower = tc.name.lower()
        if "sql" in name_lower or "注入" in name_lower:
            keywords.extend(["sql", "注入", "injection"])
        if "xss" in name_lower or "跨站" in name_lower:
            keywords.extend(["xss", "跨站", "script"])
        if "csrf" in name_lower or "跨站请求" in name_lower:
            keywords.extend(["csrf", "跨站请求伪造"])
        if "上传" in name_lower or "upload" in name_lower:
            keywords.extend(["upload", "上传", "file"])
        if "命令" in name_lower or "rce" in name_lower:
            keywords.extend(["rce", "命令执行", "command"])
        if "越权" in name_lower or "access" in name_lower:
            keywords.extend(["access", "越权", "authorization"])
        if "敏感" in name_lower or "泄露" in name_lower:
            keywords.extend(["leak", "泄露", "sensitive"])
        
        tc_keywords[tc.id] = keywords
    
    # Match each vulnerability to best test case
    used_tcs: set[str] = set()
    
    for vuln in vulnerabilities:
        vuln_title = vuln.get("title", "").lower()
        vuln_desc = vuln.get("description", "").lower()
        vuln_cwe = vuln.get("cwe", "")
        vuln_endpoint = vuln.get("endpoint", "")
        
        best_match: tuple[TestCase, float, str] | None = None
        
        for tc in test_cases:
            if tc.id in used_tcs:
                continue
            
            score = 0.0
            reasons = []
            
            # Check CWE match
            if vuln_cwe and tc.description:
                if vuln_cwe in tc.description:
                    score += 0.5
                    reasons.append(f"CWE match: {vuln_cwe}")
            
            # Check keyword match
            keywords = tc_keywords.get(tc.id, [])
            for kw in keywords:
                if kw in vuln_title or kw in vuln_desc:
                    score += 0.3
                    reasons.append(f"Keyword match: '{kw}'")
                    break
            
            # Check name contains vulnerability type
            if any(kw in vuln_title for kw in keywords):
                score += 0.2
                reasons.append("Title contains test case keywords")
            
            if score > 0:
                if best_match is None or score > best_match[1]:
                    best_match = (tc, score, "; ".join(reasons))
        
        if best_match and best_match[1] >= 0.3:
            matched_tc, confidence, reason = best_match
            matched_tc.status = "FAIL"
            matched_tc.matched_vulnerability_id = vuln.get("id")
            matched_tc.evidence = vuln.get("evidence", "")
            matched_tc.url_endpoint = vuln.get("endpoint", "") or vuln.get("target", "")
            
            matches.append(VulnerabilityMatch(
                test_case=matched_tc,
                vulnerability=vuln,
                confidence=confidence,
                match_reason=reason,
            ))
            used_tcs.add(matched_tc.id)
    
    # Mark unmatched test cases as PASS (not exploited)
    for tc in test_cases:
        if tc.id not in used_tcs:
            tc.status = "PASS"
    
    logger.info(f"Matched {len(matches)} vulnerabilities to test cases")
    return matches


def generate_compliance_report_docx(
    output_path: str | Path,
    test_cases: list[TestCase],
    matches: list[VulnerabilityMatch],
    scan_info: dict[str, Any],
) -> str:
    """Generate a DOCX compliance report following the specified template.
    
    Template structure:
    2.2 [风险等级][严重程度]测试用例名称
    - 漏洞位置：URL/API
    - 漏洞及威胁描述：预定义描述
    - 测试结果：截图和证据
    - 整改建议：预定义建议
    """
    output_path = Path(output_path)
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    
    # Title
    title = doc.add_heading('渗透测试合规报告', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Scan info
    doc.add_heading('扫描信息', level=2)
    scan_table = doc.add_table(rows=4, cols=2)
    scan_table.style = 'Table Grid'
    
    scan_data = [
        ('扫描时间', scan_info.get('scan_time', datetime.now(UTC).strftime('%Y-%m-%d %H:%M'))),
        ('目标系统', scan_info.get('target', 'N/A')),
        ('测试用例总数', str(len(test_cases))),
        ('发现漏洞数', str(len(matches))),
    ]
    
    for i, (label, value) in enumerate(scan_data):
        scan_table.cell(i, 0).text = label
        scan_table.cell(i, 1).text = value
    
    doc.add_paragraph()
    
    # Executive summary
    passed_count = sum(1 for tc in test_cases if tc.status == 'PASS')
    failed_count = len(matches)
    
    doc.add_heading('执行摘要', level=2)
    summary = (
        f"本次测试共执行 {len(test_cases)} 个测试用例，"
        f"其中 {passed_count} 个通过，{failed_count} 个发现漏洞。"
    )
    doc.add_paragraph(summary)
    
    # Detailed findings by test case
    doc.add_heading('详细测试结果', level=2)
    
    section_num = 1
    for tc in test_cases:
        # Determine risk level text
        if tc.status == 'PASS':
            risk_text = '[不存在风险]'
            severity_display = '信息'
        else:
            # Get severity from matched vulnerability
            vuln_severity = '中'
            for match in matches:
                if match.test_case.id == tc.id:
                    vuln_severity = {
                        'critical': '严重',
                        'high': '高危',
                        'medium': '中危',
                        'low': '低危',
                        'info': '信息',
                    }.get(match.vulnerability.get('severity', 'medium'), '中危')
                    break
            risk_text = f'[存在风险]'
            severity_display = vuln_severity
        
        # Section heading: 2.2[不存在风险][高危]SQL 注入漏洞检查
        section_title = f"{section_num}.{section_num}{risk_text}[{severity_display}]{tc.name}"
        doc.add_heading(section_title, level=3)
        
        # 漏洞位置
        doc.add_paragraph('漏洞位置：', style='Intense Quote')
        if tc.url_endpoint:
            doc.add_paragraph(tc.url_endpoint)
        elif tc.matched_vulnerability_id:
            # Try to extract URL from vulnerability's http_request
            for match in matches:
                if match.test_case.id == tc.id:
                    http_req = match.vulnerability.get('http_request', {})
                    if http_req and isinstance(http_req, dict):
                        url = http_req.get('url', '')
                        if url:
                            doc.add_paragraph(url)
                            break
                    # Fallback to endpoint field
                    endpoint = match.vulnerability.get('endpoint', '')
                    if endpoint:
                        doc.add_paragraph(endpoint)
                        break
            else:
                doc.add_paragraph('N/A（未针对特定端点）')
        else:
            doc.add_paragraph('N/A（未针对特定端点）')
        
        # 漏洞及威胁描述
        doc.add_paragraph('漏洞及威胁描述：', style='Intense Quote')
        if tc.status == 'FAIL' and tc.matched_vulnerability_id:
            # Use actual vulnerability description
            for match in matches:
                if match.test_case.id == tc.id:
                    vuln_desc = match.vulnerability.get('description', tc.description)
                    vuln_impact = match.vulnerability.get('impact', '')
                    doc.add_paragraph(f"{vuln_desc}\n\n影响：{vuln_impact}")
                    break
        else:
            doc.add_paragraph('经测试，该项功能未发现安全漏洞。')
        
        # 测试结果
        doc.add_paragraph('测试结果：', style='Intense Quote')
        if tc.status == 'FAIL':
            doc.add_paragraph('❌ 未通过 - 发现漏洞', style='List Bullet')
            
            # Add evidence
            if tc.evidence:
                doc.add_paragraph('证据：', style='List Bullet')
                # Clean up evidence (remove internal paths, etc.)
                evidence_text = tc.evidence[:500] + '...' if len(tc.evidence) > 500 else tc.evidence
                doc.add_paragraph(evidence_text)
            
            # Add screenshots if available
            for match in matches:
                if match.test_case.id == tc.id:
                    screenshots = match.vulnerability.get('verification_screenshots', [])
                    if screenshots:
                        doc.add_paragraph('验证截图：', style='List Bullet')
                        for ss in screenshots:
                            caption = ss.get('caption', 'Screenshot')
                            path = ss.get('path', '')
                            doc.add_paragraph(f'- {caption}: {path}', style='List Bullet')
                    break
        else:
            doc.add_paragraph('✅ 通过 - 未发现漏洞', style='List Bullet')
        
        # 整改建议
        doc.add_paragraph('整改建议：', style='Intense Quote')
        if tc.status == 'FAIL' and tc.remediation:
            # Use custom remediation from test case
            doc.add_paragraph(tc.remediation)
        elif tc.status == 'FAIL':
            # Use generic remediation based on test case name
            generic_remediation = _get_generic_remediation(tc.name)
            doc.add_paragraph(generic_remediation)
        else:
            doc.add_paragraph('无需整改。')
        
        doc.add_paragraph('_' * 50)
        section_num += 1
    
    # Save document
    doc.save(str(output_path))
    logger.info(f"Compliance report saved to {output_path}")
    return str(output_path)


def _get_generic_remediation(test_case_name: str) -> str:
    """Get generic remediation advice based on test case name."""
    name_lower = test_case_name.lower()
    
    if 'sql' in name_lower or '注入' in name_lower:
        return """1、增加对客户端提交数据的合法性验证，至少严格过滤 SQL 语句中的关键字，并且所有验证都应该在服务器端实现；
需过滤的关键字为：
[1]'单引号
[2]"双引号
[3]\\'反斜杠单引号
[4]\\"反斜杠双引号
[5]) 括号
[6]；分号
[7]--双减号
[8]+加号
[9]SQL 关键字，如 select，delete，drop 等等，注意对于关键字要对大小写都识别，如:select；SELECT；seLEcT 等都应识别；
2、建议使用较低权限的用户访问数据库。不要使用数据库管理员等高权限的用户访问数据库；
3、使用预编译技术防止 SQL 注入。"""
    elif 'xss' in name_lower or '跨站' in name_lower:
        return """1、对所有用户输入进行 HTML 实体编码；
2、实施内容安全策略 (CSP)；
3、使用 HTTPOnly 和 Secure 标志设置 Cookie。"""
    elif '上传' in name_lower or 'upload' in name_lower:
        return """1、限制允许的文件扩展名白名单；
2、验证文件 MIME 类型；
3、将上传文件存储在 Web 根目录之外；
4、对文件名进行随机化处理。"""
    else:
        return """1、遵循安全开发最佳实践；
2、定期进行安全代码审查；
3、实施纵深防御策略。"""


def create_sample_test_case_xlsx(output_path: str | Path) -> str:
    """Create a sample XLSX test case template file."""
    output_path = Path(output_path)
    wb = Workbook()
    ws = wb.active
    ws.title = "Test Cases"
    
    # Headers
    headers = ['ID', 'Name', 'Category', 'Severity', 'Description', 'Remediation']
    for col, header in enumerate(headers, start=1):
        ws.cell(row=1, column=col, value=header)
    
    # Sample test cases (Chinese)
    sample_cases = [
        [
            'TC-001',
            'SQL 注入漏洞检查',
            'Injection',
            'High',
            '检查应用程序是否存在 SQL 注入漏洞，攻击者可通过构造恶意 SQL 语句操控数据库。',
            '1、使用参数化查询或预编译语句；2、过滤特殊字符；3、最小化数据库权限。'
        ],
        [
            'TC-002',
            'XSS 跨站脚本漏洞检查',
            'Web Security',
            'Medium',
            '检查是否存在反射型或存储型 XSS 漏洞，攻击者可注入恶意脚本。',
            '1、HTML 编码输出；2、实施 CSP；3、验证和清理输入。'
        ],
        [
            'TC-003',
            '文件上传漏洞检查',
            'File Security',
            'High',
            '检查文件上传功能是否允许上传恶意文件（如 webshell）。',
            '1、白名单限制扩展名；2、验证文件内容；3、隔离存储上传文件。'
        ],
        [
            'TC-004',
            '越权访问漏洞检查',
            'Access Control',
            'High',
            '检查是否存在水平或垂直越权访问漏洞。',
            '1、实施严格的访问控制；2、服务端验证权限；3、使用基于角色的访问控制。'
        ],
        [
            'TC-005',
            '敏感信息泄露检查',
            'Data Protection',
            'Medium',
            '检查是否泄露敏感信息（如密码、密钥、个人信息等）。',
            '1、加密敏感数据；2、移除调试信息；3、实施最小化信息暴露原则。'
        ],
    ]
    
    for row_idx, case in enumerate(sample_cases, start=2):
        for col_idx, value in enumerate(case, start=1):
            ws.cell(row=row_idx, column=col_idx, value=value)
    
    wb.save(str(output_path))
    logger.info(f"Sample test case XLSX created at {output_path}")
    return str(output_path)
